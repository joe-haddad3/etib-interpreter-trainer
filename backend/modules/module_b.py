"""
Module B — TTS + Pedagogical Materials
========================================
Source: Cahier des charges, Section B — Production automatique associée au discours

Responsible: Person 4

For every generated speech, this module produces:
  B1  Audio (TTS) with configurable voice and rate
  B2  Accent variability (AR: Lebanese, Gulf, Egyptian; FR: Paris, Canadian; EN: US, UK, AU)
  B3  Full text script
  B4  Key terms list
  B5  Preparatory context / mind map (text-based)
  B6  Comprehension questions
  B7  Flashcards (term → definition, term → equivalent)
  B8  Thematic summary
  B9  MCQ (multiple-choice questions)
  B10 Key concepts identification
  B11 Editable trilingual glossary AR-FR-EN (downloadable)
  B12 Sight translation scroller interface (served by frontend)

Endpoints:
  POST /api/module-b/tts             — convert text to speech audio
  POST /api/module-b/materials       — generate all pedagogical materials for a speech
  GET  /api/module-b/audio/<filename> — serve generated audio file
"""
import os
import asyncio
import uuid
import json
import re
import io
from flask import Blueprint, request, jsonify, send_file
from config import TTS_VOICES, DEFAULT_VOICE, AUDIO_OUTPUT_FOLDER, PRIMARY_LLM_MODEL

module_b_bp = Blueprint('module_b', __name__)


# ── TTS ─────────────────────────────────────────────────────────────────────

def get_voice(language: str, accent: str = None) -> str:
    """Return the edge-tts voice name for a given language and accent."""
    lang_voices = TTS_VOICES.get(language, TTS_VOICES['en'])
    if accent and accent in lang_voices:
        return lang_voices[accent]
    return DEFAULT_VOICE.get(language, DEFAULT_VOICE['en'])

async def _tts_async(text: str, voice: str, output_path: str, rate: str = '+0%'):
    """Internal async TTS call using edge-tts."""
    import edge_tts
    communicate = edge_tts.Communicate(text, voice, rate=rate)
    await communicate.save(output_path)


_AR_D = r'[0-9٠-٩]'   # Western + Eastern Arabic-Indic digits


def _prepare_arabic_tts_text(text: str) -> str:
    """
    Make numbers read correctly by the Arabic TTS voice (student feedback,
    18 July): edge-tts reads "٨٫٢" / "8.2" as "eight two" — the decimal
    separator must be spoken as the word فاصلة so the listener hears
    "eight POINT two". Thousands separators are removed so grouped numbers
    ("١٬٤٠٠") are read as one value.
    """
    text = str(text or '')
    # Arabic thousands separator ٬ (U+066C) between digits → drop
    text = re.sub(rf'({_AR_D})٬(?={_AR_D})', r'\1', text)
    # Comma/dot as thousands group (separator + exactly 3 digits) → drop separator
    text = re.sub(rf'({_AR_D})[.,](?={_AR_D}{{3}}(?!{_AR_D}))', r'\1', text)
    # Remaining decimal separators (٫ U+066B, Arabic comma ،, dot, comma)
    # between digits → the spoken word فاصلة
    text = re.sub(rf'({_AR_D})\s*[٫،.,]\s*(?={_AR_D})', r'\1 فاصلة ', text)
    # Percent sign after a figure: edge-tts reads ٪ inconsistently in Arabic
    # (sometimes silently) — say it out loud instead. ETIB feedback 21 Aug 2026:
    # figures are among the most error-prone Arabic outputs.
    text = re.sub(rf'({_AR_D})\s*[٪%]', r'\1 في المائة', text)
    # A hyphen/dash BETWEEN two figures is a range, not a minus sign.
    text = re.sub(rf'({_AR_D})\s*[-–—]\s*(?={_AR_D})', r'\1 إلى ', text)
    return text


# ── Natural pacing ───────────────────────────────────────────────────────────
# Lina (7 Aug): the audio "feels too fast", and slowing it with the rate slider
# sounds ARTIFICIAL (edge-tts uniformly stretches the voice → robotic). The
# natural way to slow interpretation audio is a short PAUSE between sentences
# (breathing room for the interpreter), not stretching the words. We synthesise
# sentence-by-sentence and splice a small silence between them. Set the flag to
# False to revert to the original single-call synthesis.
NATURAL_PACING_ENABLED = True
_INTER_SENTENCE_PAUSE_MS = 300
_SENTENCE_SPLIT_RE = re.compile(r'(?<=[.!?؟…])\s+|\n{1,}')


def _silence_mp3_bytes(ms: int = _INTER_SENTENCE_PAUSE_MS, sr: int = 24000) -> bytes:
    """A short silent MP3 (24 kHz mono, matching edge-tts) to splice between clips."""
    import io as _io, numpy as _np, av as _av
    buf = _io.BytesIO()
    out = _av.open(buf, 'w', format='mp3')
    st = out.add_stream('libmp3lame', rate=sr)
    st.layout = 'mono'
    n = int(sr * ms / 1000)
    frame = _av.AudioFrame.from_ndarray(_np.zeros((1, n), dtype=_np.int16), format='s16', layout='mono')
    frame.sample_rate = sr
    for p in st.encode(frame):
        out.mux(p)
    for p in st.encode(None):
        out.mux(p)
    out.close()
    return buf.getvalue()


def _split_sentences(text: str) -> list:
    return [s.strip() for s in _SENTENCE_SPLIT_RE.split(str(text or '')) if s.strip()]


async def _multi_tts_bytes_async(chunks: list, voice: str, rate: str) -> list:
    return await asyncio.gather(*[_tts_bytes_async(c, voice, rate) for c in chunks])


def run_tts(text: str, language: str, accent: str = None,
            rate_adjustment: int = 0) -> str:
    """
    Convert text to speech. Returns path to the generated MP3 file.
    rate_adjustment: percentage change from normal speed (-50 to +50).
    Natural pacing: for a normal-length speech, sentences are synthesised
    separately and spliced with a short pause so the delivery is calmer without
    the artificial voice-stretch. Very long speeches (or any failure) fall back
    to the original single call.
    """
    if language == 'ar':
        text = _prepare_arabic_tts_text(text)
    voice = get_voice(language, accent)
    rate_str = f'{rate_adjustment:+d}%' if rate_adjustment != 0 else '+0%'
    filename = f'speech_{uuid.uuid4().hex[:8]}.mp3'
    output_path = os.path.join(AUDIO_OUTPUT_FOLDER, filename)

    sentences = _split_sentences(text)
    if NATURAL_PACING_ENABLED and 2 <= len(sentences) <= 25:
        try:
            parts = asyncio.run(_multi_tts_bytes_async(sentences, voice, rate_str))
            combined = _silence_mp3_bytes().join(p for p in parts if p)
            if combined:
                with open(output_path, 'wb') as fh:
                    fh.write(combined)
                return output_path, filename
        except Exception:
            pass   # any pacing failure → original single-call synthesis below

    asyncio.run(_tts_async(text, voice, output_path, rate_str))
    return output_path, filename


# ── Multi-voice (dialogue) TTS ───────────────────────────────────────────────
# Lina/team request: the 2-speaker scenarios (interview, court, medical, panel,
# broadcast) should be VOICED by two different voices in the same audio, not one
# narrator reading both parts. run_tts() above is the ORIGINAL single-voice path
# and is kept untouched. To fully revert to single-voice, set the flag below to
# False (the /tts route then always uses run_tts).
DIALOGUE_TTS_ENABLED = True

# Contrasting voices per language, alternating gender so the two speakers sound
# clearly different. Speaker order of first appearance maps to this list.
DIALOGUE_VOICES = {
    'ar': ['ar-LB-RamiNeural', 'ar-LB-LaylaNeural', 'ar-SA-HamedNeural'],
    'fr': ['fr-FR-HenriNeural', 'fr-FR-DeniseNeural', 'fr-CA-AntoineNeural'],
    'en': ['en-GB-RyanNeural', 'en-US-JennyNeural', 'en-IN-PrabhatNeural'],
}

_ROLE_WORDS = {
    'interviewer', 'guest', 'moderator', 'panelist', 'panellist', 'anchor',
    'correspondent', 'spokesperson', 'journalist', 'reporter', 'counsel',
    'witness', 'clinician', 'patient', 'doctor', 'nurse', 'host', 'officer',
    'expert', 'presenter', 'speaker', 'chair', 'delegate', 'applicant',
}

# A speaker label at the start, after a newline, or after sentence punctuation:
# 1–3 capitalised words (covers "Interviewer", "Dr. Smith", "First Last") + ":".
_LABEL_RE = re.compile(
    r'(?:^|\n|(?<=[.!?…”"]))[\s\-–—]*'
    r'([A-Z][\w.\'’\-]*(?:\s+[A-Z][\w.\'’\-]*){0,2})\s*:\s'
)


def _is_speaker_label(label: str, count: int) -> bool:
    low = label.strip().lower().rstrip('.')
    if low in _ROLE_WORDS:
        return True
    if re.match(r'^(dr|mr|mrs|ms|prof|sir|madam|hon)\.?\s', low):
        return True
    words = label.split()
    if len(words) >= 2 and all(w[:1].isupper() for w in words):   # First Last
        return True
    return count >= 2   # a recurring label is almost certainly a speaker


def _split_dialogue_turns(text: str):
    """Return [(speaker, spoken_text), ...] for a labelled dialogue, or None
    when the text is a single-speaker monologue (fewer than 2 real speakers)."""
    text = str(text or '')
    matches = list(_LABEL_RE.finditer(text))
    if len(matches) < 2:
        return None
    counts = {}
    for m in matches:
        counts[m.group(1).strip()] = counts.get(m.group(1).strip(), 0) + 1
    quals = [m for m in matches if _is_speaker_label(m.group(1).strip(), counts[m.group(1).strip()])]
    if len(quals) < 2 or len({m.group(1).strip() for m in quals}) < 2:
        return None
    turns = []
    for i, m in enumerate(quals):
        speaker = m.group(1).strip()
        start = m.end()
        end = quals[i + 1].start() if i + 1 < len(quals) else len(text)
        spoken = text[start:end].strip()
        if spoken:
            turns.append((speaker, spoken))
    # any substantial preamble before the first label → keep it (spoken by the
    # first voice) so no content is lost
    pre = text[:quals[0].start()].strip()
    if pre and len(pre.split()) >= 3 and turns:
        turns[0] = (turns[0][0], pre + ' ' + turns[0][1])
    if len(turns) < 2 or len({s for s, _ in turns}) < 2:
        return None
    return turns


# Accurate gender per voice (the config keys are NOT reliable — e.g. Lebanese
# "Rami" is male but its key "LB" has no _m suffix). Used to pick a contrasting
# second voice for two-voice dialogues.
_VOICE_GENDER = {
    'ar-LB-RamiNeural': 'm', 'ar-LB-LaylaNeural': 'f',
    'ar-SA-ZariyahNeural': 'f', 'ar-SA-HamedNeural': 'm',
    'ar-EG-SalmaNeural': 'f', 'ar-EG-ShakirNeural': 'm',
    'ar-MA-MounaNeural': 'f', 'ar-MA-JamalNeural': 'm',
    'fr-FR-DeniseNeural': 'f', 'fr-FR-HenriNeural': 'm',
    'fr-CA-SylvieNeural': 'f', 'fr-CA-AntoineNeural': 'm',
    'en-US-JennyNeural': 'f', 'en-US-GuyNeural': 'm',
    'en-GB-SoniaNeural': 'f', 'en-GB-RyanNeural': 'm',
    'en-AU-NatashaNeural': 'f', 'en-IE-EmilyNeural': 'f', 'en-IE-ConnorNeural': 'm',
    'en-IN-NeerjaNeural': 'f', 'en-IN-PrabhatNeural': 'm',
}


def _gender_of(voice_name: str) -> str:
    return _VOICE_GENDER.get(voice_name, 'f')


def _voice_base(voice_key: str) -> str:
    return re.sub(r'_(m|f)$', '', voice_key)   # 'LB'/'LB_m'/'LB_f' → 'LB'


def _dialogue_voice_list(language: str, accent: str) -> list:
    """Ordered voices for a dialogue, RESPECTING the user's accent choice:
    speaker 1 = the picked voice, speaker 2 = a contrasting voice (prefer the
    same accent, opposite gender), speaker 3 = any other distinct voice.
    Fixes: changing the accent (e.g. British → American) now changes the audio."""
    voices_map = TTS_VOICES.get(language, TTS_VOICES['en'])   # {key: voice_name}
    first_voice = get_voice(language, accent)
    first_key = accent if accent in voices_map else next(
        (k for k, v in voices_map.items() if v == first_voice), None)
    ordered, used = [first_voice], {first_voice}
    first_gender = _gender_of(first_voice)

    # 2nd: same accent, opposite gender (e.g. GB female → GB male, LB male → LB female)
    if first_key:
        for k, v in voices_map.items():
            if v not in used and _voice_base(k) == _voice_base(first_key) and _gender_of(v) != first_gender:
                ordered.append(v); used.add(v); break
    # else: any opposite-gender voice
    if len(ordered) < 2:
        for v in voices_map.values():
            if v not in used and _gender_of(v) != first_gender:
                ordered.append(v); used.add(v); break
    # last resort: any other distinct voice
    if len(ordered) < 2:
        for v in voices_map.values():
            if v not in used:
                ordered.append(v); used.add(v); break
    # 3rd distinct voice for panels (3+ speakers)
    for v in voices_map.values():
        if v not in used:
            ordered.append(v); used.add(v); break
    return ordered or DIALOGUE_VOICES.get(language, DIALOGUE_VOICES['en'])


def _assign_dialogue_voices(turns: list, language: str, accent: str = None) -> list:
    order = []
    for speaker, _ in turns:
        if speaker not in order:
            order.append(speaker)
    voices = _dialogue_voice_list(language, accent)
    mapping = {s: voices[i % len(voices)] for i, s in enumerate(order)}
    return [(mapping[s], t) for s, t in turns]


async def _tts_bytes_async(text: str, voice: str, rate: str) -> bytes:
    """Synthesize one turn and return the MP3 bytes (no temp file)."""
    import edge_tts
    communicate = edge_tts.Communicate(text, voice, rate=rate)
    buf = bytearray()
    async for chunk in communicate.stream():
        if chunk.get('type') == 'audio' and chunk.get('data'):
            buf.extend(chunk['data'])
    return bytes(buf)


async def _dialogue_tts_async(voiced_turns: list, rate: str) -> list:
    # gather preserves order → the clips concatenate in reading order
    return await asyncio.gather(*[_tts_bytes_async(t, v, rate) for v, t in voiced_turns])


def run_dialogue_tts(text: str, language: str, accent: str = None, rate_adjustment: int = 0):
    """Two-voice audio for a labelled dialogue. Returns (path, filename,
    voices_used) or None when the text is not a dialogue (caller falls back to
    single-voice run_tts). The user's accent choice picks the first speaker's voice."""
    turns = _split_dialogue_turns(text)
    if not turns:
        return None
    voiced = _assign_dialogue_voices(turns, language, accent)
    if language == 'ar':
        voiced = [(v, _prepare_arabic_tts_text(t)) for v, t in voiced]
    rate_str = f'{rate_adjustment:+d}%' if rate_adjustment != 0 else '+0%'
    parts = asyncio.run(_dialogue_tts_async(voiced, rate_str))
    # short pause between turns so the speaker changes sound natural
    try:
        combined = _silence_mp3_bytes(250).join(p for p in parts if p)
    except Exception:
        combined = b''.join(p for p in parts if p)
    if not combined:
        return None   # synthesis failed — let the caller fall back
    filename = f'speech_{uuid.uuid4().hex[:8]}.mp3'
    output_path = os.path.join(AUDIO_OUTPUT_FOLDER, filename)
    with open(output_path, 'wb') as fh:
        fh.write(combined)
    voices_used = list(dict.fromkeys(v for v, _ in voiced))
    return output_path, filename, voices_used


def _audio_duration_seconds(path: str) -> float | None:
    """Duration of a generated audio file via PyAV (already a dependency)."""
    try:
        import av
        with av.open(path) as container:
            if container.duration:
                return float(container.duration) / av.time_base
            for stream in container.streams:
                if stream.duration and stream.time_base:
                    return float(stream.duration * stream.time_base)
    except Exception:
        pass
    return None


@module_b_bp.route('/tts', methods=['POST'])
def text_to_speech():
    """
    Convert a speech script to audio.

    Request body (JSON):
      text             str   the speech text                  required
      language         str   'ar' | 'fr' | 'en'              required
      accent           str   e.g. 'LB', 'EG', 'US', 'GB'    optional
      rate_adjustment  int   -50 (slower) to +50 (faster)    default 0

    Response (JSON):
      audio_url         str    URL to retrieve the audio file
      filename          str    filename for reference
      voice_used        str    the edge-tts voice name
      duration_seconds  float  measured audio duration (when measurable)
      words_per_minute  int    measured speaking rate — professors asked for
                               an explicit wpm because simultaneous training
                               requires sources at ~100-120 wpm (Prof. LSF)
    """
    params = request.get_json()
    if not params or not params.get('text'):
        return jsonify({'error': 'text is required'}), 400

    language = params.get('language', 'ar')
    accent = params.get('accent')
    rate_adjustment = params.get('rate_adjustment', 0)
    try:
        # Two-speaker scenarios (interview/court/medical/panel/broadcast) get two
        # voices in one file; everything else falls back to single-voice run_tts.
        result = None
        if DIALOGUE_TTS_ENABLED:
            try:
                result = run_dialogue_tts(params['text'], language, accent, rate_adjustment)
            except Exception:
                result = None   # any dialogue-TTS failure → single-voice fallback

        if result:
            path, filename, voices_used = result
            is_dialogue = True
        else:
            path, filename = run_tts(text=params['text'], language=language,
                                     accent=accent, rate_adjustment=rate_adjustment)
            voices_used = [get_voice(language, accent)]
            is_dialogue = False

        duration = _audio_duration_seconds(path)
        word_count = len(str(params['text']).split())
        wpm = round(word_count / (duration / 60.0)) if duration and duration > 1 else None
        return jsonify({
            'audio_url': f'/api/module-b/audio/{filename}',
            'filename': filename,
            'voice_used': voices_used[0],
            'voices_used': voices_used,
            'dialogue': is_dialogue,
            'duration_seconds': round(duration, 1) if duration else None,
            'words_per_minute': wpm,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@module_b_bp.route('/audio/<filename>')
def serve_audio(filename: str):
    """Serve a generated audio file (basename only — no path traversal)."""
    safe_name = os.path.basename(filename)
    if not safe_name.endswith('.mp3'):
        return jsonify({'error': 'File not found'}), 404
    path = os.path.join(AUDIO_OUTPUT_FOLDER, safe_name)
    if not os.path.exists(path):
        return jsonify({'error': 'File not found'}), 404
    return send_file(path, mimetype='audio/mpeg')


# ── Pedagogical materials ────────────────────────────────────────────────────

MATERIALS_PROMPT = """You are a pedagogical materials developer for ETIB (École de Traducteurs et d'Interprètes de Beyrouth, USJ Beirut).

Given the following {lang_name} conference speech on the topic of "{domain}", generate training materials for interpreter trainees.

SPEECH:
{script}

Return ONLY a valid JSON object — no markdown, no code fences, no explanation. Use this exact structure:

{{
  "key_terms": ["term1", "term2", "term3", "term4", "term5"],
  "summary": "[Main Theme]\\n└── Context: brief background\\n└── Pillar 1\\n    ├── Sub-point A\\n    └── Key data point\\n└── Pillar 2\\n    └── Sub-point B",
  "mcq": [
    {{"question": "Question text?", "options": ["A. option1", "B. option2", "C. option3", "D. option4"], "answer": "A"}},
    {{"question": "Question text?", "options": ["A. option1", "B. option2", "C. option3", "D. option4"], "answer": "C"}}
  ],
  "comprehension": ["Open question 1?", "Open question 2?"],
  "glossary": [
    {{"ar": "Arabic term", "fr": "French term", "en": "English term"}},
    {{"ar": "Arabic term", "fr": "French term", "en": "English term"}},
    {{"ar": "Arabic term", "fr": "French term", "en": "English term"}},
    {{"ar": "Arabic term", "fr": "French term", "en": "English term"}},
    {{"ar": "Arabic term", "fr": "French term", "en": "English term"}}
  ]
}}

Rules:
- key_terms: exactly 5 important domain-specific terms from the speech (in the speech language)
- summary: visual text tree — keep it concise, max 10 lines
- mcq: 2 to 3 questions with 4 options each, mark the correct answer letter
- glossary: exactly 5 entries, every entry must have Arabic, French AND English
"""


def _extract_json(text: str) -> dict:
    """Extract JSON from LLM response, stripping markdown code fences if present."""
    text = text.strip()
    fence = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
    if fence:
        text = fence.group(1).strip()
    start = text.find('{')
    end = text.rfind('}') + 1
    if start >= 0 and end > start:
        text = text[start:end]
    return json.loads(text)


@module_b_bp.route('/materials', methods=['POST'])
def generate_materials():
    """
    Generate all pedagogical materials for a given speech script.

    Request body (JSON):
      script     str   the speech text                  required
      language   str   'ar' | 'fr' | 'en'              default 'ar'
      domain     str   speech domain                    default 'general'

    Response (JSON):
      key_terms       list[str]
      summary         str   (mind-map formatted text)
      mcq             list[{question, options, answer}]
      comprehension   list[str]
      glossary        list[{ar, fr, en}]
    """
    data = request.get_json()
    if not data or not data.get('script'):
        return jsonify({'error': 'script is required'}), 400

    script   = data['script']
    language = data.get('language', 'ar')
    domain   = data.get('domain', 'general')

    lang_names = {'ar': 'Arabic', 'fr': 'French', 'en': 'English'}

    try:
        from utils.groq_client import get_groq_client
        client = get_groq_client()

        prompt = MATERIALS_PROMPT.format(
            lang_name=lang_names.get(language, 'English'),
            domain=domain,
            script=script
        )

        response = client.chat.completions.create(
            model=PRIMARY_LLM_MODEL,
            messages=[
                {
                    'role': 'system',
                    'content': (
                        'You are a pedagogical materials developer for interpreter training. '
                        'You always return valid JSON only — no markdown, no extra text.'
                    )
                },
                {'role': 'user', 'content': prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )

        content = response.choices[0].message.content
        materials = _extract_json(content)
        return jsonify(materials)

    except json.JSONDecodeError as e:
        return jsonify({'error': f'Failed to parse LLM response: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@module_b_bp.route('/glossary/download', methods=['POST'])
def download_glossary():
    """
    Export a trilingual glossary as a downloadable DOCX file.
    Requirement: Cahier des charges B11 — editable and downloadable.

    Request body (JSON):
      glossary   list[{ar, fr, en}]   required
      domain     str                  used in filename
    """
    data = request.get_json()
    if not data or not data.get('glossary'):
        return jsonify({'error': 'glossary is required'}), 400

    glossary = data['glossary']
    domain   = data.get('domain', 'General')

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        section = doc.sections[0]
        section.left_margin  = Inches(1)
        section.right_margin = Inches(1)

        title = doc.add_heading('ETIB — Trilingual Glossary', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f'Domain: {domain}')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ['العربية / Arabic', 'Français / French', 'English']):
            cell.text = label
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(11)

        for entry in glossary:
            row = table.add_row().cells
            row[0].text = entry.get('ar', '')
            row[1].text = entry.get('fr', '')
            row[2].text = entry.get('en', '')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', domain)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'glossary_{safe_name}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ── Glossary upload / import ─────────────────────────────────────────────────
# Professor bilan (23 July): "Glossaire à éditer avec possibilité de le
# téléverser dans la plateforme pour qu'il soit pris en compte dans la phase
# d'évaluation." The student uploads their own trilingual glossary; it is parsed
# into the standard {term, arabic, french, english, definition} shape and merged
# into the in-app glossary, which already flows into Module D as approved
# terminology.

_GLOSSARY_COLUMN_ALIASES = {
    'term':       ('term', 'terme', 'مصطلح', 'key term', 'source term', 'concept'),
    'arabic':     ('arabic', 'arabe', 'العربية', 'عربي', 'ar', 'arabic term', 'المقابل العربي'),
    'french':     ('french', 'français', 'francais', 'fr', 'الفرنسية', 'terme français'),
    'english':    ('english', 'anglais', 'en', 'الإنجليزية', 'english term'),
    'definition': ('definition', 'définition', 'definitions', 'meaning', 'تعريف', 'def'),
}


def _glossary_field_for_header(header: str) -> str | None:
    """
    Map an arbitrary column header to a canonical glossary field name.

    Must handle BILINGUAL headers — our own DOCX export writes
    "العربية / Arabic", "Français / French", "English" (tester bug Aug 2026:
    re-importing a platform-exported glossary only filled the English column
    because exact matching failed on the combined headers).
    """
    key = str(header or '').strip().casefold()
    if not key:
        return None
    # 1. Exact match on the whole header
    for field, aliases in _GLOSSARY_COLUMN_ALIASES.items():
        if key in aliases:
            return field
    # 2. Token match — split compound headers ("العربية / Arabic") and try
    #    each part. First matching token wins.
    tokens = [t for t in re.split(r'[/|,;()\[\]–—-]+|\s+', key) if t]
    for tok in tokens:
        for field, aliases in _GLOSSARY_COLUMN_ALIASES.items():
            if tok in aliases:
                return field
    # 3. Substring match for long aliases only (≥4 chars — keeps 'ar'/'fr'/'en'
    #    from false-matching inside unrelated words).
    for field, aliases in _GLOSSARY_COLUMN_ALIASES.items():
        for alias in aliases:
            if len(alias) >= 4 and alias in key:
                return field
    return None


def _normalize_uploaded_glossary_rows(rows: list[dict]) -> list[dict]:
    """Keep only real glossary fields; drop rows that carry no content."""
    cleaned = []
    for row in rows:
        entry = {k: str(row.get(k, '') or '').strip() for k in
                 ('term', 'arabic', 'french', 'english', 'definition')}
        # A row needs at least one language equivalent or a term to be useful.
        if entry['term'] or entry['arabic'] or entry['french'] or entry['english']:
            cleaned.append(entry)
    return cleaned


def _parse_glossary_json(data: bytes) -> list[dict]:
    parsed = json.loads(data.decode('utf-8-sig'))
    if isinstance(parsed, dict):
        parsed = parsed.get('glossary', [])
    if not isinstance(parsed, list):
        return []
    rows = []
    for item in parsed:
        if not isinstance(item, dict):
            continue
        row = {}
        for header, value in item.items():
            field = _glossary_field_for_header(header)
            if field:
                row[field] = value
        rows.append(row)
    return _normalize_uploaded_glossary_rows(rows)


def _parse_glossary_csv(data: bytes) -> list[dict]:
    import csv
    text = data.decode('utf-8-sig', errors='replace')
    # Sniff the delimiter (comma / semicolon / tab) from the first line.
    first_line = text.splitlines()[0] if text.strip() else ''
    delimiter = ';' if first_line.count(';') > first_line.count(',') else ','
    if '\t' in first_line and first_line.count('\t') >= first_line.count(delimiter):
        delimiter = '\t'
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    table = [r for r in reader if any(cell.strip() for cell in r)]
    if not table:
        return []
    header_map = {i: _glossary_field_for_header(h) for i, h in enumerate(table[0])}
    has_header = any(field for field in header_map.values())
    rows = []
    if has_header:
        for raw in table[1:]:
            row = {}
            for i, cell in enumerate(raw):
                field = header_map.get(i)
                if field:
                    row[field] = cell
            rows.append(row)
    else:
        # No recognizable header — assume columns: term, arabic, french, english, definition
        order = ('term', 'arabic', 'french', 'english', 'definition')
        for raw in table:
            rows.append({order[i]: cell for i, cell in enumerate(raw) if i < len(order)})
    return _normalize_uploaded_glossary_rows(rows)


def _parse_glossary_docx(data: bytes) -> list[dict]:
    from docx import Document
    doc = Document(io.BytesIO(data))
    rows = []
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text for c in table.rows[0].cells]
        header_map = {i: _glossary_field_for_header(h) for i, h in enumerate(header_cells)}
        has_header = any(field for field in header_map.values())
        body_rows = table.rows[1:] if has_header else table.rows
        order = ('arabic', 'french', 'english')   # matches the DOCX export column order
        for trow in body_rows:
            cells = [c.text for c in trow.cells]
            row = {}
            if has_header:
                for i, cell in enumerate(cells):
                    field = header_map.get(i)
                    if field:
                        row[field] = cell
            else:
                for i, cell in enumerate(cells):
                    if i < len(order):
                        row[order[i]] = cell
            rows.append(row)
    return _normalize_uploaded_glossary_rows(rows)


def _parse_glossary_txt(data: bytes) -> list[dict]:
    text = data.decode('utf-8-sig', errors='replace')
    rows = []
    for line in text.splitlines():
        if not line.strip():
            continue
        # Accept "term | arabic | french | english | definition" or tab-separated.
        parts = re.split(r'\s*\|\s*|\t', line.strip())
        parts = [p.strip() for p in parts if p.strip()]
        if not parts:
            continue
        order = ('term', 'arabic', 'french', 'english', 'definition')
        rows.append({order[i]: parts[i] for i in range(min(len(parts), len(order)))})
    return _normalize_uploaded_glossary_rows(rows)


@module_b_bp.route('/glossary/upload', methods=['POST'])
def upload_glossary():
    """
    Parse an uploaded glossary file (CSV, TSV, TXT, DOCX, or JSON) into the
    standard trilingual glossary shape so it can be merged into the in-app
    glossary and used by the evaluation phase. Returns {'glossary': [...]}.
    """
    file = request.files.get('glossary') or request.files.get('file')
    if not file or not file.filename:
        return jsonify({'error': 'No glossary file provided'}), 400

    filename = file.filename.lower()
    raw = file.read()
    if not raw:
        return jsonify({'error': 'The uploaded file is empty'}), 400
    if len(raw) > 2 * 1024 * 1024:
        return jsonify({'error': 'Glossary file too large (max 2 MB)'}), 400

    try:
        if filename.endswith('.json'):
            entries = _parse_glossary_json(raw)
        elif filename.endswith(('.csv', '.tsv')):
            entries = _parse_glossary_csv(raw)
        elif filename.endswith('.docx'):
            entries = _parse_glossary_docx(raw)
        elif filename.endswith(('.txt', '.md')):
            entries = _parse_glossary_txt(raw)
        else:
            return jsonify({'error': 'Unsupported file type. Use CSV, TSV, TXT, DOCX, or JSON.'}), 400
    except Exception as e:
        return jsonify({'error': f'Could not read the glossary file: {e}'}), 400

    if not entries:
        return jsonify({'error': 'No glossary terms could be read from the file. '
                                 'Expected columns like Term / Arabic / French / English.'}), 400

    return jsonify({'glossary': entries[:200]})


# ── Per-user glossary correction memory ──────────────────────────────────────
# Lina (7 Aug 2026): "can the platform remember, PER USER, the terms it
# corrects, so the glossary improves over time?" Stored in MongoDB keyed by
# user_id + normalized term, so a logged-in user's corrections follow their
# account across sessions/devices (guests have no server memory — local only).
_glossary_mongo = None
_glossary_mem: dict = {}   # { user_id: { term_norm: {term, arabic, french, english, definition} } }
_GLOSSARY_MONGODB_URI = os.getenv('MONGODB_URI', 'mongodb://127.0.0.1:27017')
_GLOSSARY_MONGODB_DB = os.getenv('MONGODB_DB', 'etib_interpreter_trainer')
_GLOSSARY_FIELDS = ('arabic', 'french', 'english', 'definition')


def _glossary_user_id() -> str:
    """User id from the auth token, or '' for guests/anonymous."""
    try:
        from modules.auth import get_user_from_token
        token = (
            request.args.get('auth_token', '') or
            request.form.get('auth_token', '') or
            (request.get_json(silent=True) or {}).get('auth_token', '')
        ).strip()
        user = get_user_from_token(token)
        return user['id'] if user else ''
    except Exception:
        return ''


def _glossary_corrections_collection():
    global _glossary_mongo
    if _glossary_mongo is not None:
        return _glossary_mongo
    try:
        from pymongo import MongoClient, ASCENDING
        client = MongoClient(_GLOSSARY_MONGODB_URI, serverSelectionTimeoutMS=2000)
        client.admin.command('ping')
        col = client[_GLOSSARY_MONGODB_DB]['glossary_corrections']
        col.create_index([('user_id', ASCENDING), ('term_norm', ASCENDING)], unique=True)
        _glossary_mongo = col
    except Exception:
        _glossary_mongo = None
    return _glossary_mongo


def _norm_glossary_term(term: str) -> str:
    return str(term or '').strip().lower()


@module_b_bp.route('/glossary/corrections', methods=['GET'])
def get_glossary_corrections():
    """Return every glossary correction saved by the logged-in user."""
    user_id = _glossary_user_id()
    if not user_id:
        return jsonify({'corrections': []})   # guests: no server memory
    col = _glossary_corrections_collection()
    if col is not None:
        docs = list(col.find({'user_id': user_id}, {'_id': 0, 'user_id': 0, 'updated_at': 0}))
    else:
        docs = list(_glossary_mem.get(user_id, {}).values())
    return jsonify({'corrections': docs})


@module_b_bp.route('/glossary/corrections', methods=['POST'])
def save_glossary_corrections():
    """Upsert one or more glossary corrections for the logged-in user.
    Accepts {'correction': {...}} or {'corrections': [{...}, ...]}. Each item
    needs a 'term' plus at least one of arabic/french/english/definition."""
    user_id = _glossary_user_id()
    if not user_id:
        return jsonify({'error': 'login required'}), 401
    data = request.get_json(silent=True) or {}
    items = data.get('corrections')
    if items is None and isinstance(data.get('correction'), dict):
        items = [data['correction']]
    if not isinstance(items, list):
        return jsonify({'error': 'corrections must be a list or a single correction'}), 400

    import datetime
    col = _glossary_corrections_collection()
    saved = 0
    for item in items:
        if not isinstance(item, dict):
            continue
        term = str(item.get('term', '')).strip()
        term_norm = _norm_glossary_term(term)
        if not term_norm:
            continue
        record = {'term': term, 'term_norm': term_norm}
        for field in _GLOSSARY_FIELDS:
            value = item.get(field)
            if isinstance(value, str) and value.strip():
                record[field] = value.strip()
        if not any(field in record for field in _GLOSSARY_FIELDS):
            continue   # nothing worth storing (blank correction)
        if col is not None:
            col.update_one(
                {'user_id': user_id, 'term_norm': term_norm},
                {'$set': {**record, 'user_id': user_id, 'updated_at': datetime.datetime.utcnow()}},
                upsert=True,
            )
        else:
            _glossary_mem.setdefault(user_id, {})[term_norm] = record
        saved += 1
    return jsonify({'saved': saved})

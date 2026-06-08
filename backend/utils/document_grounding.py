"""
Phase 1 document grounding helpers for Module A.

This module intentionally uses only local text extraction, chunking, and
keyword overlap scoring.
"""
import os
import re
from io import BytesIO


SUPPORTED_DOCUMENT_EXTENSIONS = {'.txt', '.docx', '.pdf'}
MIN_EXTRACTED_CHARACTERS = 120
DEFAULT_CHUNK_CHARACTERS = 1800
DEFAULT_CHUNK_OVERLAP = 250
DEFAULT_MAX_EXCERPTS = 4
DEFAULT_MAX_EXCERPT_CHARACTERS = 6500
RETRIEVAL_METHOD = 'keyword_overlap'


class DocumentGroundingError(ValueError):
    """User-facing document grounding error."""


def get_source_type(filename: str) -> str:
    """Return the normalized file extension, including the leading dot."""
    return os.path.splitext(filename or '')[1].lower()


def is_supported_document(filename: str) -> bool:
    """Return True if the filename has a supported document extension."""
    return get_source_type(filename) in SUPPORTED_DOCUMENT_EXTENSIONS


def extract_document_text(file_storage, source_type: str) -> str:
    """Extract text from a Flask FileStorage for a supported source type."""
    file_storage.stream.seek(0)
    data = file_storage.read()
    file_storage.stream.seek(0)

    if source_type == '.txt':
        return _extract_txt(data)
    if source_type == '.docx':
        return _extract_docx(data)
    if source_type == '.pdf':
        return _extract_pdf(data)

    raise DocumentGroundingError(
        'Unsupported document type. Please upload a .txt, .docx, or .pdf file.'
    )


def normalize_text(text: str) -> str:
    """Normalize extracted document text into compact readable text."""
    text = text or ''
    text = text.replace('\x00', ' ')
    text = re.sub(r'[ \t\r\f\v]+', ' ', text)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    text = re.sub(r' *\n *', '\n', text)
    return text.strip()


def validate_extracted_text(text: str) -> None:
    """Reject documents that do not provide enough usable text."""
    if len(text) < MIN_EXTRACTED_CHARACTERS:
        raise DocumentGroundingError(
            'The uploaded document does not contain enough readable text. '
            'Please upload a longer text-based PDF, DOCX, or TXT file.'
        )


def chunk_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_CHARACTERS,
    overlap: int = DEFAULT_CHUNK_OVERLAP
) -> list[str]:
    """Split normalized text into overlapping character chunks."""
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)
        if end < text_length:
            boundary = max(text.rfind('\n\n', start, end), text.rfind('. ', start, end))
            if boundary > start + int(chunk_size * 0.55):
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_length:
            break
        start = max(0, end - overlap)

    return chunks


def select_relevant_chunks(
    chunks: list[str],
    params: dict,
    max_excerpts: int = DEFAULT_MAX_EXCERPTS,
    max_total_characters: int = DEFAULT_MAX_EXCERPT_CHARACTERS
) -> list[str]:
    """Select chunks with simple metadata/keyword overlap scoring."""
    if not chunks:
        return []

    query_terms = _metadata_terms(params)
    scored = []

    for index, chunk in enumerate(chunks):
        chunk_terms = _tokenize(chunk)
        score = sum(1 for term in query_terms if term in chunk_terms)

        if params.get('number_density') == 'high':
            score += min(5, len(re.findall(r'\d+', chunk)))
        elif params.get('number_density') == 'medium':
            score += min(3, len(re.findall(r'\d+', chunk)))

        score += max(0, 3 - index) * 0.1
        scored.append((score, index, chunk))

    scored.sort(key=lambda item: (-item[0], item[1]))

    selected = []
    total_chars = 0
    for _, _, chunk in scored:
        if len(selected) >= max_excerpts:
            break
        if total_chars and total_chars + len(chunk) > max_total_characters:
            continue
        selected.append(chunk)
        total_chars += len(chunk)

    if not selected:
        selected = [chunks[0][:max_total_characters]]

    return selected


def score_chunks(chunk_records: list[dict], params: dict) -> list[dict]:
    """Score chunk records with keyword and metadata overlap."""
    query_terms = _tokenize(str(params.get('query') or ''))
    metadata_terms = _metadata_terms(params)
    scored_chunks = []

    for record in chunk_records:
        text = record.get('text', '')
        chunk_terms = _tokenize(text)
        query_score = sum(3 for term in query_terms if term in chunk_terms)
        metadata_score = sum(1 for term in metadata_terms if term in chunk_terms)
        score = query_score + metadata_score

        if params.get('number_density') == 'high':
            score += min(5, len(re.findall(r'\d+', text)))
        elif params.get('number_density') == 'medium':
            score += min(3, len(re.findall(r'\d+', text)))

        chunk_index = record.get('chunk_index', 0)
        score += max(0, 3 - chunk_index) * 0.1

        scored_record = dict(record)
        scored_record['score'] = score
        scored_record['retrieval_method'] = RETRIEVAL_METHOD
        scored_chunks.append(scored_record)

    return scored_chunks


def select_relevant_chunks_with_metadata(
    chunk_records: list[dict],
    params: dict,
    max_excerpts: int = DEFAULT_MAX_EXCERPTS,
    max_total_characters: int = DEFAULT_MAX_EXCERPT_CHARACTERS,
) -> list[dict]:
    """Select relevant chunk records while preserving source metadata."""
    scored_chunks = score_chunks(chunk_records, params)
    scored_chunks.sort(key=lambda item: (
        -item.get('score', 0),
        item.get('source_order', 0),
        item.get('chunk_index', 0)
    ))

    selected = []
    total_chars = 0
    for chunk in scored_chunks:
        if len(selected) >= max_excerpts:
            break

        text_length = len(chunk.get('text', ''))
        if total_chars and total_chars + text_length > max_total_characters:
            continue

        selected.append({
            'text': chunk.get('text', ''),
            'source_filename': chunk.get('source_filename'),
            'source_type': chunk.get('source_type'),
            'chunk_index': chunk.get('chunk_index', 0),
            'score': chunk.get('score', 0),
            'retrieval_method': chunk.get('retrieval_method', RETRIEVAL_METHOD),
        })
        total_chars += text_length

    return selected


def format_excerpts_for_prompt(excerpts: list[str]) -> str:
    """Format selected excerpts for the LLM prompt."""
    return '\n\n'.join(
        f'[Excerpt {index}]\n{excerpt}'
        for index, excerpt in enumerate(excerpts, start=1)
    )


def _extract_txt(data: bytes) -> str:
    for encoding in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentGroundingError('Could not decode the uploaded TXT file.')


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentGroundingError(
            'DOCX support requires python-docx to be installed.'
        ) from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise DocumentGroundingError(
            'Could not read the uploaded DOCX file. Please check the file and try again.'
        ) from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    table_cells.append(cell.text)

    return '\n'.join(paragraphs + table_cells)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentGroundingError(
            'PDF support requires pypdf to be installed. Run pip install -r requirements.txt.'
        ) from exc

    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or '' for page in reader.pages]
    except Exception as exc:
        raise DocumentGroundingError(
            'Could not read the uploaded PDF file. Please use a text-based PDF.'
        ) from exc

    return '\n'.join(pages)


def _metadata_terms(params: dict) -> set[str]:
    values = [
        params.get('query'),
        params.get('domain'),
        params.get('scenario'),
        params.get('difficulty'),
        params.get('mode'),
        params.get('number_density'),
        params.get('language'),
    ]
    terms = set()
    for value in values:
        terms.update(_tokenize(str(value or '')))

    language = params.get('language')
    language_names = {
        'ar': ['arabic', 'arab', 'modern', 'standard'],
        'fr': ['french', 'francais', 'france'],
        'en': ['english'],
    }
    for term in language_names.get(language, []):
        terms.add(term)

    return {term for term in terms if len(term) > 2}


def _tokenize(text: str) -> set[str]:
    return {token.lower() for token in re.findall(r'\w+', text or '', flags=re.UNICODE)}
